"""
Renderer - Converts resume JSON to Word DOCX
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, Any
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


class Renderer:
    """Renders resume JSON to formatted Word DOCX"""
    
    GITHUB_URLS = {
        "Dino Game Deep RL Agent": "https://github.com/virtual457/dino-game-AI",
        "Orion Platform": "https://github.com/virtual457/Orion-platform",
        "Orion PaaS": "https://github.com/virtual457/Orion-platform",
        "Kambaz Learning Management System": "https://github.com/virtual457/kambaz-next-js",
        "Calendly": "https://github.com/virtual457/Calendly",
        "Port Management System": "https://github.com/virtual457/Port-Management-System",
        "Maritime Logistics Platform": "https://github.com/virtual457/Port-Management-System",
        "Data Analysis on PUBG": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Large Scale Data Analysis": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Face Recognition & Validation System": "https://github.com/virtual457/Recognition-and-Validation-of-Faces-using-Machine-Learning-and-Image-Processing",
        "Online Examination System": "https://github.com/virtual457/Online-examination-using-mongodb",
        "Calendly - Calendar Management System": "https://github.com/virtual457/Calendly"
    }
    
    def __init__(self, template_path: str):
        """Initialize with template path"""
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
    
    def render(self, resume_json: Dict[str, Any], output_path: str) -> str:
        """
        Render resume JSON to DOCX
        
        Args:
            resume_json: Resume data
            output_path: Where to save DOCX
        
        Returns:
            Path to saved file
        """
        doc = Document(self.template_path)
        
        self._update_header(doc, resume_json)
        self._update_summary(doc, resume_json.get('summary', ''))
        self._update_skills(doc, resume_json.get('skills', []))
        self._update_experience(doc, resume_json.get('experience', []))
        self._update_projects(doc, resume_json.get('projects', []))
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return str(output_path.absolute())
    
    def _add_text_with_bold_markers(self, paragraph, text: str, font_size: int = 10):
        """Add text with **bold** markers"""
        parts = text.split('**')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.size = Pt(font_size)
            if i % 2 == 1:  # Odd indices are bold
                run.bold = True
    
    def _add_hyperlink(self, paragraph, text: str, url: str, font_size: int = 10, bold: bool = False):
        """Add hyperlink to paragraph"""
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size * 2))
            rPr.append(sz)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._element.append(hyperlink)
    
    def _find_paragraph_by_text(self, doc: Document, search_text: str) -> int:
        """Find paragraph index containing text"""
        for i, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                return i
        return None
    
    def _update_header(self, doc: Document, resume_json: Dict[str, Any]):
        """Update header"""
        if len(doc.paragraphs) < 3:
            return
        
        doc.paragraphs[0].text = "Chandan Gowda K S"
        
        title = resume_json.get('header', {}).get('title', 'Software Engineer | MS CS @ Northeastern')
        doc.paragraphs[1].text = title
        
        contact_para = doc.paragraphs[2]
        contact_para.clear()
        
        run = contact_para.add_run("+1 (857) 421-7469; ")
        run.font.size = Pt(10)
        
        self._add_hyperlink(contact_para, "chandan.keelara@gmail.com", "mailto:chandan.keelara@gmail.com", 10)
        contact_para.add_run("; ")
        
        self._add_hyperlink(contact_para, "LinkedIn", "https://www.linkedin.com/in/chandan-gowda-k-s-765194186/", 10)
        contact_para.add_run("; ")
        
        self._add_hyperlink(contact_para, "Portfolio", "https://virtual457.github.io/", 10)
        contact_para.add_run("; ")
        
        self._add_hyperlink(contact_para, "GitHub", "https://github.com/virtual457", 10)
        contact_para.add_run(";")
    
    def _update_summary(self, doc: Document, summary_text: str):
        """Update summary"""
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith("Software Engineer"):
                para.clear()
                self._add_text_with_bold_markers(para, summary_text, 10)
                return
    
    def _update_skills(self, doc: Document, skills_list: list):
        """Update skills section"""
        skills_idx = self._find_paragraph_by_text(doc, "TECHNICAL SKILLS")
        if skills_idx is None:
            return
        
        tab_position = 2.45
        
        for i, skill in enumerate(skills_list[:7]):
            para_idx = skills_idx + 1 + i
            if para_idx >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[para_idx]
            para.clear()
            
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(tab_position), WD_TAB_ALIGNMENT.LEFT)
            
            category_run = para.add_run(skill.get('category', ''))
            category_run.bold = True
            category_run.font.size = Pt(10)
            
            para.add_run("\t")
            
            items_run = para.add_run(skill.get('items', ''))
            items_run.bold = False
            items_run.font.size = Pt(10)
    
    def _update_experience(self, doc: Document, experience_list: list):
        """Update work experience"""
        exp_idx = self._find_paragraph_by_text(doc, "WORK EXPERIENCE")
        if exp_idx is None:
            return
        
        current_para = exp_idx + 1
        
        for company_exp in experience_list:
            current_para += 1  # Skip company header
            
            bullets = company_exp.get('bullets', [])
            for bullet_text in bullets:
                if current_para >= len(doc.paragraphs):
                    break
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, bullet_text, 10)
                current_para += 1
    
    def _update_projects(self, doc: Document, projects_list: list):
        """Update projects section"""
        projects_idx = self._find_paragraph_by_text(doc, "PROJECTS")
        if projects_idx is None:
            return
        
        current_para = projects_idx + 1
        
        for project in projects_list[:3]:
            if current_para >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[current_para]
            para.clear()
            
            project_title = project.get('title', '')
            github_url = self.GITHUB_URLS.get(project_title)
            
            if github_url:
                self._add_hyperlink(para, project_title, github_url, 10, bold=True)
            else:
                title_run = para.add_run(project_title)
                title_run.bold = True
                title_run.font.size = Pt(10)
            
            para.add_run(" | ")
            
            tech_run = para.add_run(project.get('tech', ''))
            tech_run.bold = True
            tech_run.italic = True
            tech_run.font.size = Pt(10)
            
            para.add_run(" | ")
            
            if github_url:
                self._add_hyperlink(para, "GitHub", github_url, 10)
            else:
                github_run = para.add_run("GitHub")
                github_run.font.size = Pt(10)
            
            current_para += 1
            
            # Bullet 1
            if current_para < len(doc.paragraphs):
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, project.get('bullet1', ''), 10)
                current_para += 1
            
            # Bullet 2
            if current_para < len(doc.paragraphs):
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, project.get('bullet2', ''), 10)
                current_para += 1
