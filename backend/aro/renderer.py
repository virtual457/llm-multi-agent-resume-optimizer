"""
Renderer - Converts resume JSON to Word DOCX

This module takes resume JSON from the Generator Agent and produces
a formatted Word document WITHOUT opening Microsoft Word application.

Pure python-docx manipulation for server compatibility.

Usage:
    renderer = ResumeRenderer("path/to/template.docx")
    renderer.render(resume_json, "output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, Any, List
import os


class ResumeRenderer:
    """Renders resume JSON to formatted Word DOCX"""
    
    # GitHub URLs mapping
    GITHUB_URLS = {
        "Dino Game Deep RL Agent": "https://github.com/virtual457/dino-game-AI",
        "Orion PaaS": "https://github.com/virtual457/Orion-platform",
        "Orion Platform": "https://github.com/virtual457/Orion-platform",
        "Calendly - Calendar Management System": "https://github.com/virtual457/Calendly",
        "Calendly": "https://github.com/virtual457/Calendly",
        "Maritime Logistics Platform": "https://github.com/virtual457/Port-Management-System",
        "Port Management System": "https://github.com/virtual457/Port-Management-System",
        "Large Scale Data Analysis": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Data Analysis on PUBG": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Face Recognition & Validation System": "https://github.com/virtual457/Recognition-and-Validation-of-Faces-using-Machine-Learning-and-Image-Processing",
        "Online Examination System": "https://github.com/virtual457/Online-examination-using-mongodb",
        "Kambaz Learning Management System": "https://github.com/virtual457/kambaz-next-js"
    }
    
    def __init__(self, template_path: str):
        """
        Initialize renderer with template
        
        Args:
            template_path: Path to Word template file
        """
        self.template_path = Path(template_path)
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        print(f"📄 Renderer initialized with template: {self.template_path}")
    
    def render(self, resume_json: Dict[str, Any], output_path: str) -> str:
        """
        Render resume JSON to Word DOCX
        
        Args:
            resume_json: Resume data from Generator
            output_path: Where to save the output file
        
        Returns:
            Path to generated file
        """
        print("\n" + "="*60)
        print("RESUME RENDERER - JSON → DOCX")
        print("="*60 + "\n")
        
        # Load template
        print("📂 Loading Word template...")
        doc = Document(self.template_path)
        print(f"   ✅ Template loaded ({len(doc.paragraphs)} paragraphs)\n")
        
        # Update each section
        self._update_header(doc, resume_json)
        self._update_summary(doc, resume_json.get('summary', ''))
        self._update_skills(doc, resume_json.get('skills', []))
        self._update_experience(doc, resume_json.get('experience', []))
        self._update_projects(doc, resume_json.get('projects', []))
        
        # Save output
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        print("\n💾 Saving generated resume...")
        doc.save(str(output_path))
        print(f"   ✅ Saved to: {output_path}\n")
        
        print("="*60)
        print("✅ RENDERING COMPLETE!")
        print("="*60)
        print(f"\n📄 Output file: {output_path.absolute()}\n")
        
        return str(output_path.absolute())
    
    def _add_text_with_bold_markers(self, paragraph, text: str, font_size: int = 10, base_bold: bool = False):
        """
        Add text with **bold** markers
        
        Example: "Built **AWS Lambda** pipeline" → "Built AWS Lambda pipeline" 
                 (with AWS Lambda bold)
        
        Args:
            paragraph: Word paragraph object
            text: Text with **markers** for bold sections
            font_size: Font size for all text
            base_bold: If True, non-marked text is also bold
        """
        parts = text.split('**')
        
        for i, part in enumerate(parts):
            if not part:  # Skip empty parts
                continue
            
            run = paragraph.add_run(part)
            run.font.size = Pt(font_size)
            
            # Odd indices (1, 3, 5...) = text between ** markers = bold
            if i % 2 == 1:
                run.bold = True
            else:
                run.bold = base_bold
    
    def _add_hyperlink(self, paragraph, text: str, url: str, font_size: int = None, bold: bool = False):
        """
        Add a hyperlink to a paragraph
        
        Args:
            paragraph: Word paragraph object
            text: Display text
            url: URL to link to
            font_size: Font size (points)
            bold: Whether to make text bold
        """
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(
            qn('r:id'), 
            paragraph.part.relate_to(
                url, 
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
        )
        
        # Create run for hyperlink
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add bold if specified
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        
        # Add underline and color (blue) for hyperlink style
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Blue
        rPr.append(color)
        
        # Add font size if specified
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size * 2))  # Word uses half-points
            rPr.append(sz)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._element.append(hyperlink)
    
    def _find_paragraph_by_text(self, doc: Document, search_text: str) -> int:
        """Find paragraph index containing search text"""
        for i, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                return i
        return None
    
    def _update_header(self, doc: Document, resume_json: Dict[str, Any]):
        """Update header (first 3 paragraphs: name, title, contact)"""
        print("📝 Updating header...")
        
        if len(doc.paragraphs) < 3:
            print("   ⚠️  Warning: Not enough paragraphs for header")
            return
        
        # Line 1: Name
        doc.paragraphs[0].text = "Chandan Gowda K S"
        
        # Line 2: Title (from JSON or generate)
        title = resume_json.get('header', {}).get('title', 'Software Engineer | MS CS @ Northeastern')
        doc.paragraphs[1].text = title
        
        # Line 3: Contact info with hyperlinks
        contact_para = doc.paragraphs[2]
        contact_para.clear()
        
        # Phone
        run = contact_para.add_run("+1 (857) 421-7469; ")
        run.font.size = Pt(10)
        
        # Email
        self._add_hyperlink(
            contact_para, 
            "chandan.keelara@gmail.com", 
            "mailto:chandan.keelara@gmail.com", 
            font_size=10
        )
        contact_para.add_run("; ")
        
        # LinkedIn
        self._add_hyperlink(
            contact_para,
            "LinkedIn",
            "https://www.linkedin.com/in/chandan-gowda-k-s-765194186/",
            font_size=10
        )
        contact_para.add_run("; ")
        
        # Portfolio
        self._add_hyperlink(
            contact_para,
            "Portfolio",
            "https://virtual457.github.io/",
            font_size=10
        )
        contact_para.add_run("; ")
        
        # GitHub
        self._add_hyperlink(
            contact_para,
            "GitHub",
            "https://github.com/virtual457",
            font_size=10
        )
        contact_para.add_run(";")
        
        print("   ✅ Header updated (with hyperlinks)")
    
    def _update_summary(self, doc: Document, summary_text: str):
        """Update summary paragraph"""
        print("📝 Updating summary...")
        
        # Find summary paragraph (starts with "Software Engineer")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith("Software Engineer"):
                para.clear()
                self._add_text_with_bold_markers(para, summary_text, font_size=10)
                print("   ✅ Summary updated (with bold markers)")
                return
        
        print("   ⚠️  Warning: Could not find summary paragraph")
    
    def _update_skills(self, doc: Document, skills_list: List[Dict[str, str]]):
        """Update TECHNICAL SKILLS section"""
        print("📝 Updating technical skills...")
        
        # Find TECHNICAL SKILLS heading
        skills_idx = self._find_paragraph_by_text(doc, "TECHNICAL SKILLS")
        
        if skills_idx is None:
            print("   ⚠️  Warning: Could not find TECHNICAL SKILLS section")
            return
        
        # Fixed tab position at 35 characters (2.45 inches)
        tab_position = 2.45
        
        # Update next 7 paragraphs
        for i, skill in enumerate(skills_list[:7]):  # Max 7 categories
            para_idx = skills_idx + 1 + i
            if para_idx >= len(doc.paragraphs):
                print(f"   ⚠️  Warning: Not enough paragraphs for skill {i+1}")
                break
            
            para = doc.paragraphs[para_idx]
            para.clear()
            
            # Set tab stop
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(tab_position), WD_TAB_ALIGNMENT.LEFT)
            
            # Add category (bold)
            category_run = para.add_run(skill.get('category', ''))
            category_run.bold = True
            category_run.font.size = Pt(10)
            
            # Add tab
            para.add_run("\t")
            
            # Add items (not bold)
            items_run = para.add_run(skill.get('items', ''))
            items_run.bold = False
            items_run.font.size = Pt(10)
        
        print(f"   ✅ Updated {len(skills_list)} skill categories")
    
    def _update_experience(self, doc: Document, experience_list: List[Dict[str, Any]]):
        """Update WORK EXPERIENCE section"""
        print("📝 Updating work experience...")
        
        # Find WORK EXPERIENCE heading
        exp_idx = self._find_paragraph_by_text(doc, "WORK EXPERIENCE")
        
        if exp_idx is None:
            print("   ⚠️  Warning: Could not find WORK EXPERIENCE section")
            return
        
        current_para = exp_idx + 1
        
        for company_exp in experience_list:
            # Skip company header line (stays in template)
            current_para += 1
            
            # Update bullets
            bullets = company_exp.get('bullets', [])
            
            for bullet_text in bullets:
                if current_para >= len(doc.paragraphs):
                    print(f"   ⚠️  Warning: Not enough paragraphs for bullets")
                    break
                
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, bullet_text, font_size=10)
                current_para += 1
            
            company_name = company_exp.get('company', 'Company')
            print(f"   ✅ Updated {len(bullets)} bullets for {company_name}")
        
        print("   ✅ Work experience section complete")
    
    def _update_projects(self, doc: Document, projects_list: List[Dict[str, Any]]):
        """Update PROJECTS section with GitHub hyperlinks"""
        print("📝 Updating projects...")
        
        # Find PROJECTS heading
        projects_idx = self._find_paragraph_by_text(doc, "PROJECTS")
        
        if projects_idx is None:
            print("   ⚠️  Warning: Could not find PROJECTS section")
            return
        
        current_para = projects_idx + 1
        
        for project in projects_list[:3]:  # Max 3 projects
            if current_para >= len(doc.paragraphs):
                print(f"   ⚠️  Warning: Not enough paragraphs for projects")
                break
            
            # Update title line
            para = doc.paragraphs[current_para]
            para.clear()
            
            # Add project title as hyperlink (BOLD)
            project_title = project.get('title', 'Project')
            github_url = self.GITHUB_URLS.get(project_title)
            
            if github_url:
                self._add_hyperlink(para, project_title, github_url, font_size=10, bold=True)
            else:
                title_run = para.add_run(project_title)
                title_run.bold = True
                title_run.font.size = Pt(10)
            
            para.add_run(" | ")
            
            # Add tech stack (BOLD ITALIC)
            tech_run = para.add_run(project.get('tech', ''))
            tech_run.bold = True
            tech_run.italic = True
            tech_run.font.size = Pt(10)
            
            para.add_run(" | ")
            
            # Add GitHub link
            if github_url:
                self._add_hyperlink(para, "GitHub", github_url, font_size=10)
            else:
                github_run = para.add_run("GitHub")
                github_run.font.size = Pt(10)
            
            current_para += 1
            
            # Update bullet 1
            if current_para < len(doc.paragraphs):
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, project.get('bullet1', ''), font_size=10)
                current_para += 1
            
            # Update bullet 2
            if current_para < len(doc.paragraphs):
                para = doc.paragraphs[current_para]
                para.clear()
                self._add_text_with_bold_markers(para, project.get('bullet2', ''), font_size=10)
                current_para += 1
        
        print(f"   ✅ Updated {len(projects_list)} projects (with GitHub hyperlinks)")


def test_renderer():
    """Test the renderer with sample JSON"""
    import json
    
    print("\n" + "="*60)
    print("RENDERER TEST")
    print("="*60 + "\n")
    
    # Sample resume JSON (simplified)
    sample_json = {
        "header": {
            "title": "Backend Engineer | MS CS @ Northeastern | Python, AWS, Kubernetes"
        },
        "summary": "MS Computer Science student at Northeastern (**3.89 GPA**) with **4+ years** production experience building **distributed systems** at **LSEG** and **Infosys**. Specialized in **Python**, **AWS Lambda/SQS**, **Kubernetes**, and **machine learning** inference systems. Delivered **7.5M+ record** processing pipelines with **40% latency** reduction. Excited to apply scalable backend expertise to solve complex infrastructure challenges.",
        "skills": [
            {"category": "Languages", "items": "Python, Java, Go, JavaScript, TypeScript, SQL, C++"},
            {"category": "Cloud & DevOps", "items": "AWS (Lambda, SQS, EC2, S3), Kubernetes, Docker, CI/CD, Terraform"},
            {"category": "Backend", "items": "Microservices, REST APIs, Event-Driven Architecture, Distributed Systems"},
            {"category": "Databases", "items": "MySQL, PostgreSQL, MongoDB, Redis, DynamoDB"},
            {"category": "ML/AI", "items": "PyTorch, TensorFlow, Deep RL, Neural Networks, LLM Integration"},
            {"category": "Frameworks", "items": "Django, Flask, FastAPI, Spring Boot, Node.js, Express"},
            {"category": "Tools", "items": "Git, Linux, Bash, Prometheus, Grafana, Kafka"}
        ],
        "experience": [
            {
                "company": "London Stock Exchange Group (LSEG)",
                "role": "Senior Software Engineer",
                "location": "Bengaluru",
                "duration": "08-2022 to 12-2024",
                "bullets": [
                    "Built **distributed Python/Java pipelines** processing **7.5M+ records** across **180+ countries**, integrating **AWS Lambda**, **SQS**, and **microservices** to automate financial data workflows with **40% latency reduction**",
                    "Engineered **event-driven architecture** using **AWS SQS** and **Lambda** for async processing, handling **100K+ daily messages** with **99.9% reliability** and automated error recovery",
                    "Developed **REST APIs** serving **50+ internal services**, implementing **OAuth 2.0** authentication and **rate limiting** for secure, scalable access to financial datasets",
                    "Optimized **SQL queries** and **database schemas** in **MySQL/PostgreSQL**, reducing query times by **60%** through indexing strategies and query plan analysis",
                    "Led migration to **Kubernetes** for containerized deployments, implementing **CI/CD pipelines** with **Jenkins** and **Docker**, cutting deployment time from hours to minutes"
                ]
            },
            {
                "company": "Infosys",
                "role": "Senior Systems Engineer",
                "location": "Bengaluru",
                "duration": "10-2020 to 07-2022",
                "bullets": [
                    "Built **Python data pipelines** processing **500GB+ daily** from **REST APIs** and **SQL databases**, achieving **3x throughput improvement** through parallel processing and caching",
                    "Developed **ETL workflows** with **Apache Airflow** for automated data ingestion, transformation, and loading into **data warehouses**",
                    "Implemented **monitoring dashboards** using **Prometheus** and **Grafana**, tracking pipeline health, latency metrics, and system performance in real-time",
                    "Collaborated with **cross-functional teams** to deliver scalable solutions, participating in code reviews and **Agile sprints**"
                ]
            }
        ],
        "projects": [
            {
                "title": "Dino Game Deep RL Agent",
                "tech": "PyTorch, Deep Q-Networks, OpenCV, Python",
                "bullet1": "Built **deep reinforcement learning agent** using **PyTorch** with **1.5M+ parameters**, achieving **90% accuracy** in gameplay through **DQN algorithm** and experience replay optimization",
                "bullet2": "Engineered real-time **inference pipeline** processing **game frames at 16.67 FPS** with **neural network** predictions, implementing epsilon-greedy exploration for continuous learning"
            },
            {
                "title": "Orion Platform",
                "tech": "Go, Kubernetes Operator SDK, Custom Resources, Docker",
                "bullet1": "Developed **custom Kubernetes operator** in **Go** managing **application lifecycles** through **CRDs**, automating deployments, scaling, and health monitoring across clusters",
                "bullet2": "Implemented **reconciliation loops** and **controller logic** for resource management, handling **100+ deployments** with automated rollback and self-healing capabilities"
            },
            {
                "title": "Port Management System",
                "tech": "Django, MySQL, Dijkstra Algorithm, REST APIs",
                "bullet1": "Built **full-stack maritime logistics platform** with **Django/MySQL**, implementing **Dijkstra's shortest path** algorithm in **SQL stored procedures** for optimal route planning",
                "bullet2": "Designed **50+ stored procedures** for complex queries, achieving **sub-second response times** for route calculations across **10K+ port combinations**"
            }
        ]
    }
    
    # Find template
    template_path = Path(__file__).parent.parent.parent / "templates" / "Chandan_Resume_Format.docx"
    
    if not template_path.exists():
        print(f"❌ Template not found at: {template_path}")
        print("\n💡 Please copy template from:")
        print("   D:\\Git\\virtual457-projects\\job-application-automator\\templates\\Chandan_Resume_Format.docx")
        print("   to:")
        print(f"   {template_path}")
        return
    
    # Create renderer
    renderer = ResumeRenderer(str(template_path))
    
    # Render resume
    output_path = Path(__file__).parent.parent.parent / "output" / "Test_Generated_Resume.docx"
    
    try:
        result = renderer.render(sample_json, str(output_path))
        print(f"\n🎉 SUCCESS! Resume generated at:\n   {result}")
        
        # Save JSON for reference
        json_path = output_path.with_suffix('.json')
        with open(json_path, 'w') as f:
            json.dump(sample_json, f, indent=2)
        print(f"\n📄 Sample JSON saved to:\n   {json_path}")
        
    except Exception as e:
        print(f"\n❌ Error during rendering: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    test_renderer()
